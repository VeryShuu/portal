"""KB export builders: filesystem-safe names and Markdown→PDF/DOCX rendering.

Pure / IO-only logic with no HTTP layer; the API handlers add ACL and audit.
"""

from __future__ import annotations

import re


def article_md_stem(title: str) -> str:
    """Filesystem-safe stem for a single-article Markdown export."""
    return re.sub(r"[^\w\- ]", "", title)[:60].strip() or "article"


def section_zip_stem(title: str) -> str:
    """Filesystem-safe stem for a section ZIP export."""
    return re.sub(r"[^\w\- ]", "", title)[:40] or "section"


def document_stem(title: str) -> str:
    """Filesystem-safe stem for PDF/DOCX document exports."""
    return re.sub(r"[^\w\s-]", "", title)[:80].strip() or "article"


def render_article_docx(title: str, body: str) -> bytes:
    """Render an article to DOCX bytes."""
    import io

    import markdown_it
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=0)

    md = markdown_it.MarkdownIt()
    tokens = md.parse(body)
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "heading_open":
            level = int(tok.tag[1])
            content_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            heading_text = content_tok.content if content_tok else ""
            doc.add_heading(heading_text, level=min(level, 9))
            i += 3
        elif tok.type == "paragraph_open":
            content_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            para_text = content_tok.content if content_tok else ""
            doc.add_paragraph(para_text)
            i += 3
        elif tok.type == "fence":
            p = doc.add_paragraph()
            run = p.add_run(tok.content)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            i += 1
        elif tok.type == "bullet_list_open" or tok.type in ("list_item_open",):
            i += 1
        elif tok.type == "inline" and i > 0 and tokens[i - 1].type == "list_item_open":
            doc.add_paragraph(tok.content, style="List Bullet")
            i += 1
        else:
            i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
