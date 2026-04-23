"""Knowledge Base API: разделы, статьи, версии, теги, комментарии, правки, обратная связь."""
from __future__ import annotations

import re
import uuid
from datetime import datetime, timezone
from typing import Any

from fastapi import APIRouter, HTTPException, Query, status
from fastapi.responses import Response
from sqlalchemy import delete, func, select, text, update
from sqlalchemy.orm import selectinload

from app.api.deps import AdminDep, CurrentUser, DbDep, EditorDep, RedisDep
from app.core.logging import get_logger
from app.core.sanitize import sanitize_html
from app.services.kb_acl import (
    require_article_permission,
    require_section_permission,
    resolve_article_permission,
    resolve_section_permission,
)
from app.models.kb import (
    KbArticle,
    KbArticleComment,
    KbArticleFeedback,
    KbArticleTag,
    KbArticleVersion,
    KbSection,
    KbSuggestion,
    KbTag,
)
from app.models.user import User
from app.schemas.kb import (
    CreateArticleRequest,
    CreateCommentRequest,
    CreateSectionRequest,
    CreateSuggestionRequest,
    DraftSaveRequest,
    FeedbackRequest,
    FeedbackStats,
    KbArticleList,
    KbArticleListItem,
    KbArticlePublic,
    KbBreadcrumb,
    KbCommentList,
    KbCommentPublic,
    KbSectionPublic,
    KbSuggestionPublic,
    KbTagPublic,
    KbUserRef,
    KbVersionList,
    KbVersionPublic,
    ReviewSuggestionRequest,
    UpdateArticleRequest,
    UpdateSectionRequest,
)
from app.services.audit import push_audit_event

router = APIRouter(prefix="/kb", tags=["knowledge-base"])
logger = get_logger(__name__)

VIEW_DEDUP_TTL = 3600  # 1 час


# ── Вспомогательные функции ───────────────────────────────────────────────────

def _slugify(text_: str) -> str:
    slug = text_.lower().strip()
    slug = re.sub(r"[^\w\s-]", "", slug, flags=re.UNICODE)
    slug = re.sub(r"[\s_-]+", "-", slug)
    slug = re.sub(r"^-+|-+$", "", slug)
    return slug or "section"


async def _get_breadcrumbs(db: Any, section_id: uuid.UUID | None) -> list[KbBreadcrumb]:
    if not section_id:
        return []
    result = await db.execute(
        text("""
            WITH RECURSIVE crumbs AS (
                SELECT id, parent_id, title, slug, 0 AS depth
                FROM kb_sections WHERE id = :section_id
                UNION ALL
                SELECT s.id, s.parent_id, s.title, s.slug, c.depth + 1
                FROM kb_sections s
                JOIN crumbs c ON s.id = c.parent_id
                WHERE c.depth < 10
            )
            SELECT id, title, slug FROM crumbs ORDER BY depth DESC
        """),
        {"section_id": str(section_id)},
    )
    rows = result.fetchall()
    return [KbBreadcrumb(id=r[0], title=r[1], slug=r[2]) for r in rows]


async def _get_article_or_404(db: Any, article_id: uuid.UUID) -> KbArticle:
    result = await db.execute(
        select(KbArticle)
        .options(selectinload(KbArticle.tags))
        .where(KbArticle.id == article_id, KbArticle.deleted_at.is_(None))
    )
    article = result.scalar_one_or_none()
    if not article:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Article not found")
    return article


def _article_to_public(article: KbArticle, breadcrumbs: list[KbBreadcrumb],
                        creator: User | None, updater: User | None,
                        helpful: int = 0, not_helpful: int = 0,
                        user_feedback: bool | None = None) -> KbArticlePublic:
    return KbArticlePublic(
        id=article.id,
        title=article.title,
        body=article.body,
        section_id=article.section_id,
        status=article.status,
        version=article.version,
        view_count=article.view_count,
        published_at=article.published_at,
        created_at=article.created_at,
        updated_at=article.updated_at,
        tags=[KbTagPublic(id=t.id, name=t.name, slug=t.slug) for t in (article.tags or [])],
        breadcrumbs=breadcrumbs,
        created_by=KbUserRef(id=creator.id, full_name=creator.full_name, avatar_url=creator.avatar_url) if creator else None,
        updated_by=KbUserRef(id=updater.id, full_name=updater.full_name, avatar_url=updater.avatar_url) if updater else None,
        helpful_count=helpful,
        not_helpful_count=not_helpful,
        user_feedback=user_feedback,
    )


async def _resolve_tags(db: Any, tag_names: list[str]) -> list[KbTag]:
    tags: list[KbTag] = []
    for name in tag_names:
        slug = _slugify(name)
        result = await db.execute(select(KbTag).where(KbTag.slug == slug))
        tag = result.scalar_one_or_none()
        if not tag:
            tag = KbTag(name=name.strip(), slug=slug)
            db.add(tag)
            await db.flush()
        tags.append(tag)
    return tags


async def _set_article_tags(db: Any, article: KbArticle, tag_names: list[str]) -> None:
    await db.execute(delete(KbArticleTag).where(KbArticleTag.article_id == article.id))
    tags = await _resolve_tags(db, tag_names)
    for tag in tags:
        db.add(KbArticleTag(article_id=article.id, tag_id=tag.id))


# ── Разделы ───────────────────────────────────────────────────────────────────

@router.get("/sections", summary="Дерево разделов")
async def get_sections(db: DbDep, user: CurrentUser, redis: RedisDep) -> dict:
    result = await db.execute(select(KbSection).order_by(KbSection.sort_order, KbSection.title))
    sections = result.scalars().all()

    section_map: dict[uuid.UUID, KbSectionPublic] = {}
    for s in sections:
        perm = await resolve_section_permission(user, s, db, redis)
        if perm is None:
            continue
        section_map[s.id] = KbSectionPublic(
            id=s.id,
            parent_id=s.parent_id,
            title=s.title,
            slug=s.slug,
            description=s.description,
            sort_order=s.sort_order,
            created_at=s.created_at,
            children=[],
        )

    roots: list[KbSectionPublic] = []
    for s in sections:
        if s.id not in section_map:
            continue
        node = section_map[s.id]
        if s.parent_id and s.parent_id in section_map:
            section_map[s.parent_id].children.append(node)
        else:
            roots.append(node)

    return {"items": roots}


@router.post("/sections", status_code=status.HTTP_201_CREATED, summary="Создать раздел")
async def create_section(
    body: CreateSectionRequest,
    db: DbDep,
    user: EditorDep,
) -> KbSectionPublic:
    slug = _slugify(body.title)
    result = await db.execute(select(KbSection).where(KbSection.slug == slug))
    if result.scalar_one_or_none():
        slug = f"{slug}-{uuid.uuid4().hex[:6]}"

    section = KbSection(
        title=body.title,
        slug=slug,
        parent_id=body.parent_id,
        description=body.description,
        sort_order=body.sort_order,
        created_by=user.id,
    )
    db.add(section)
    await db.commit()
    await db.refresh(section)
    return KbSectionPublic(
        id=section.id,
        parent_id=section.parent_id,
        title=section.title,
        slug=section.slug,
        description=section.description,
        sort_order=section.sort_order,
        created_at=section.created_at,
        children=[],
    )


@router.put("/sections/{section_id}", summary="Обновить раздел")
async def update_section(
    section_id: uuid.UUID,
    body: UpdateSectionRequest,
    db: DbDep,
    user: EditorDep,
) -> KbSectionPublic:
    result = await db.execute(select(KbSection).where(KbSection.id == section_id))
    section = result.scalar_one_or_none()
    if not section:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Section not found")

    if body.title is not None:
        section.title = body.title
    if body.description is not None:
        section.description = body.description
    if body.sort_order is not None:
        section.sort_order = body.sort_order
    if body.parent_id is not None:
        if body.parent_id == section_id:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Section cannot be its own parent")
        section.parent_id = body.parent_id

    await db.commit()
    await db.refresh(section)
    return KbSectionPublic(
        id=section.id,
        parent_id=section.parent_id,
        title=section.title,
        slug=section.slug,
        description=section.description,
        sort_order=section.sort_order,
        created_at=section.created_at,
        children=[],
    )


@router.delete("/sections/{section_id}", status_code=status.HTTP_204_NO_CONTENT, summary="Удалить раздел")
async def delete_section(
    section_id: uuid.UUID,
    db: DbDep,
    user: AdminDep,
    redis: RedisDep,
    force: bool = Query(default=False),
) -> None:
    result = await db.execute(select(KbSection).where(KbSection.id == section_id))
    section = result.scalar_one_or_none()
    if not section:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Section not found")

    child_result = await db.execute(select(KbSection).where(KbSection.parent_id == section_id))
    has_children = child_result.scalar_one_or_none() is not None
    article_result = await db.execute(select(KbArticle).where(KbArticle.section_id == section_id, KbArticle.deleted_at.is_(None)))
    has_articles = article_result.scalar_one_or_none() is not None

    if (has_children or has_articles) and not force:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Раздел содержит дочерние элементы. Используйте ?force=true",
        )

    await db.delete(section)
    await db.commit()
    await push_audit_event(
        redis,
        event_type="kb.section_deleted",
        user_id=str(user.id),
        user_email=user.email,
        resource_type="kb_section",
        resource_id=str(section_id),
    )


# ── Статьи ────────────────────────────────────────────────────────────────────

@router.get("/articles", response_model=KbArticleList, summary="Список статей KB")
async def list_articles(
    db: DbDep,
    user: CurrentUser,
    redis: RedisDep,
    section_id: uuid.UUID | None = Query(default=None),
    tag: str | None = Query(default=None),
    status_filter: str | None = Query(default=None, alias="status"),
    q: str | None = Query(default=None),
    limit: int = Query(default=20, ge=1, le=100),
    offset: int = Query(default=0, ge=0),
) -> KbArticleList:
    if status_filter in ("draft", "archived") and user.role not in ("editor", "admin"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Insufficient permissions")

    stmt = (
        select(KbArticle)
        .options(selectinload(KbArticle.tags))
        .where(KbArticle.deleted_at.is_(None))
        .order_by(KbArticle.updated_at.desc())
    )

    if not status_filter:
        if user.role in ("editor", "admin"):
            pass
        else:
            stmt = stmt.where(KbArticle.status == "published")
    else:
        stmt = stmt.where(KbArticle.status == status_filter)

    if section_id:
        stmt = stmt.where(KbArticle.section_id == section_id)

    if tag:
        tag_result = await db.execute(select(KbTag).where(KbTag.slug == _slugify(tag)))
        tag_obj = tag_result.scalar_one_or_none()
        if tag_obj:
            stmt = stmt.join(KbArticleTag, KbArticleTag.article_id == KbArticle.id).where(
                KbArticleTag.tag_id == tag_obj.id
            )
        else:
            return KbArticleList(items=[], total=0, limit=limit, offset=offset)

    if q:
        stmt = stmt.where(
            KbArticle.body_tsvector.op("@@")(func.plainto_tsquery("russian_hunspell", q))
        )

    count_stmt = select(func.count()).select_from(stmt.subquery())
    total = (await db.execute(count_stmt)).scalar_one()

    result = await db.execute(stmt.limit(limit).offset(offset))
    articles = result.scalars().all()

    creators: dict[uuid.UUID, User] = {}
    creator_ids = {a.created_by for a in articles if a.created_by}
    if creator_ids:
        u_result = await db.execute(select(User).where(User.id.in_(creator_ids)))
        for u in u_result.scalars():
            creators[u.id] = u

    items = []
    for a in articles:
        perm = await resolve_article_permission(user, a, db, redis)
        if perm is None:
            continue
        creator = creators.get(a.created_by) if a.created_by else None
        items.append(KbArticleListItem(
            id=a.id,
            title=a.title,
            section_id=a.section_id,
            status=a.status,
            version=a.version,
            view_count=a.view_count,
            published_at=a.published_at,
            created_at=a.created_at,
            updated_at=a.updated_at,
            tags=[KbTagPublic(id=t.id, name=t.name, slug=t.slug) for t in (a.tags or [])],
            created_by=KbUserRef(id=creator.id, full_name=creator.full_name, avatar_url=creator.avatar_url) if creator else None,
        ))

    return KbArticleList(items=items, total=total, limit=limit, offset=offset)


@router.post("/articles", status_code=status.HTTP_201_CREATED, response_model=KbArticlePublic, summary="Создать статью")
async def create_article(
    body: CreateArticleRequest,
    db: DbDep,
    user: EditorDep,
    redis: RedisDep,
) -> KbArticlePublic:
    if body.status not in ("draft", "published"):
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Invalid status")

    article = KbArticle(
        section_id=body.section_id,
        title=sanitize_html(body.title) if body.title else body.title,
        body=body.body,
        status=body.status,
        version=1,
        created_by=user.id,
        updated_by=user.id,
        published_at=datetime.now(timezone.utc) if body.status == "published" else None,
    )
    db.add(article)
    await db.flush()

    if body.tags:
        await _set_article_tags(db, article, body.tags)

    await db.commit()
    await db.refresh(article)

    breadcrumbs = await _get_breadcrumbs(db, article.section_id)
    await push_audit_event(
        redis,
        event_type="kb.article_created",
        user_id=str(user.id),
        user_email=user.email,
        resource_type="kb_article",
        resource_id=str(article.id),
        resource_title=article.title,
    )
    return _article_to_public(article, breadcrumbs, user, user)


@router.get("/articles/{article_id}", response_model=KbArticlePublic, summary="Получить статью")
async def get_article(
    article_id: uuid.UUID,
    db: DbDep,
    user: CurrentUser,
    redis: RedisDep,
) -> KbArticlePublic:
    article = await _get_article_or_404(db, article_id)

    await require_article_permission(user, article, "viewer", db, redis)

    if article.status != "published" and user.role not in ("editor", "admin"):
        perm = await resolve_article_permission(user, article, db, redis)
        if perm not in ("editor", "manager"):
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

    view_key = f"kb:view:{article_id}:{user.id}"
    if not await redis.get(view_key):
        await db.execute(
            update(KbArticle).where(KbArticle.id == article_id).values(view_count=KbArticle.view_count + 1)
        )
        await db.commit()
        await redis.setex(view_key, VIEW_DEDUP_TTL, "1")
        article.view_count += 1

    creator = updater = None
    if article.created_by:
        r = await db.execute(select(User).where(User.id == article.created_by))
        creator = r.scalar_one_or_none()
    if article.updated_by and article.updated_by != article.created_by:
        r = await db.execute(select(User).where(User.id == article.updated_by))
        updater = r.scalar_one_or_none()
    if not updater:
        updater = creator

    helpful_r = await db.execute(
        select(func.count()).where(KbArticleFeedback.article_id == article_id, KbArticleFeedback.is_helpful.is_(True))
    )
    not_helpful_r = await db.execute(
        select(func.count()).where(KbArticleFeedback.article_id == article_id, KbArticleFeedback.is_helpful.is_(False))
    )
    user_fb_r = await db.execute(
        select(KbArticleFeedback.is_helpful).where(
            KbArticleFeedback.article_id == article_id, KbArticleFeedback.user_id == user.id
        )
    )
    user_feedback = user_fb_r.scalar_one_or_none()

    breadcrumbs = await _get_breadcrumbs(db, article.section_id)
    return _article_to_public(
        article, breadcrumbs, creator, updater,
        helpful=helpful_r.scalar_one(),
        not_helpful=not_helpful_r.scalar_one(),
        user_feedback=user_feedback,
    )


@router.put("/articles/{article_id}", response_model=KbArticlePublic, summary="Обновить статью")
async def update_article(
    article_id: uuid.UUID,
    body: UpdateArticleRequest,
    db: DbDep,
    user: EditorDep,
    redis: RedisDep,
) -> KbArticlePublic:
    article = await _get_article_or_404(db, article_id)

    await require_article_permission(user, article, "editor", db, redis)

    if article.version != body.version:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail=f"Статья изменена другим пользователем",
            headers={"X-Current-Version": str(article.version), "X-Your-Version": str(body.version)},
        )

    version_snapshot = KbArticleVersion(
        article_id=article.id,
        version=article.version,
        title=article.title,
        body=article.body,
        changed_by=user.id,
        change_comment=body.change_comment,
    )
    db.add(version_snapshot)

    if body.title is not None:
        article.title = body.title
    if body.body is not None:
        article.body = body.body
    if body.section_id is not None:
        article.section_id = body.section_id
    if body.status is not None:
        if body.status not in ("draft", "published", "archived"):
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Invalid status")
        if body.status == "published" and article.published_at is None:
            article.published_at = datetime.now(timezone.utc)
        article.status = body.status
    if body.tags is not None:
        await _set_article_tags(db, article, body.tags)

    article.version += 1
    article.updated_by = user.id
    article.updated_at = datetime.now(timezone.utc)

    await db.commit()
    await db.refresh(article)

    breadcrumbs = await _get_breadcrumbs(db, article.section_id)
    creator = updater = None
    if article.created_by:
        r = await db.execute(select(User).where(User.id == article.created_by))
        creator = r.scalar_one_or_none()
    await push_audit_event(
        redis,
        event_type="kb.article_updated",
        user_id=str(user.id),
        user_email=user.email,
        resource_type="kb_article",
        resource_id=str(article.id),
        resource_title=article.title,
    )
    return _article_to_public(article, breadcrumbs, creator, user)


@router.put("/articles/{article_id}/draft", response_model=KbArticlePublic, summary="Автосохранение черновика")
async def save_draft(
    article_id: uuid.UUID,
    body: DraftSaveRequest,
    db: DbDep,
    user: EditorDep,
) -> KbArticlePublic:
    article = await _get_article_or_404(db, article_id)
    if article.status != "draft":
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Only drafts can be auto-saved this way")

    if body.title is not None:
        article.title = body.title
    if body.body is not None:
        article.body = body.body
    article.updated_at = datetime.now(timezone.utc)
    article.updated_by = user.id

    await db.commit()
    await db.refresh(article)

    breadcrumbs = await _get_breadcrumbs(db, article.section_id)
    creator = None
    if article.created_by:
        r = await db.execute(select(User).where(User.id == article.created_by))
        creator = r.scalar_one_or_none()
    return _article_to_public(article, breadcrumbs, creator, user)


@router.delete("/articles/{article_id}", status_code=status.HTTP_204_NO_CONTENT, summary="Удалить статью (soft)")
async def delete_article(
    article_id: uuid.UUID,
    db: DbDep,
    user: AdminDep,
    redis: RedisDep,
) -> None:
    article = await _get_article_or_404(db, article_id)
    await require_article_permission(user, article, "manager", db, redis)
    article.deleted_at = datetime.now(timezone.utc)
    await db.commit()
    await push_audit_event(
        redis,
        event_type="kb.article_deleted",
        user_id=str(user.id),
        user_email=user.email,
        resource_type="kb_article",
        resource_id=str(article_id),
    )


@router.post("/articles/{article_id}/restore", response_model=KbArticlePublic, summary="Восстановить статью")
async def restore_article(
    article_id: uuid.UUID,
    db: DbDep,
    user: AdminDep,
) -> KbArticlePublic:
    result = await db.execute(
        select(KbArticle).options(selectinload(KbArticle.tags)).where(KbArticle.id == article_id)
    )
    article = result.scalar_one_or_none()
    if not article:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Article not found")
    article.deleted_at = None
    await db.commit()
    await db.refresh(article)
    breadcrumbs = await _get_breadcrumbs(db, article.section_id)
    return _article_to_public(article, breadcrumbs, None, None)


# ── Версии ────────────────────────────────────────────────────────────────────

@router.get("/articles/{article_id}/versions", response_model=KbVersionList, summary="Версии статьи")
async def list_versions(
    article_id: uuid.UUID,
    db: DbDep,
    user: CurrentUser,
    limit: int = Query(default=20, ge=1, le=100),
    offset: int = Query(default=0, ge=0),
) -> KbVersionList:
    await _get_article_or_404(db, article_id)

    count_r = await db.execute(
        select(func.count()).where(KbArticleVersion.article_id == article_id)
    )
    total = count_r.scalar_one()

    result = await db.execute(
        select(KbArticleVersion)
        .where(KbArticleVersion.article_id == article_id)
        .order_by(KbArticleVersion.version.desc())
        .limit(limit).offset(offset)
    )
    versions = result.scalars().all()

    user_ids = {v.changed_by for v in versions if v.changed_by}
    users_map: dict[uuid.UUID, User] = {}
    if user_ids:
        u_r = await db.execute(select(User).where(User.id.in_(user_ids)))
        for u in u_r.scalars():
            users_map[u.id] = u

    items = []
    for v in versions:
        changer = users_map.get(v.changed_by) if v.changed_by else None
        items.append(KbVersionPublic(
            id=v.id, article_id=v.article_id, version=v.version,
            title=v.title, body=v.body, change_comment=v.change_comment,
            changed_by=KbUserRef(id=changer.id, full_name=changer.full_name, avatar_url=changer.avatar_url) if changer else None,
            created_at=v.created_at,
        ))

    return KbVersionList(items=items, total=total)


@router.post("/articles/{article_id}/versions/{version_number}/restore",
             response_model=KbArticlePublic, summary="Откат к версии N")
async def restore_version(
    article_id: uuid.UUID,
    version_number: int,
    db: DbDep,
    user: EditorDep,
) -> KbArticlePublic:
    article = await _get_article_or_404(db, article_id)

    v_result = await db.execute(
        select(KbArticleVersion).where(
            KbArticleVersion.article_id == article_id,
            KbArticleVersion.version == version_number,
        )
    )
    version_snap = v_result.scalar_one_or_none()
    if not version_snap:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Version not found")

    snapshot = KbArticleVersion(
        article_id=article.id,
        version=article.version,
        title=article.title,
        body=article.body,
        changed_by=user.id,
        change_comment=f"Откат к версии {version_number}",
    )
    db.add(snapshot)

    article.title = version_snap.title or article.title
    article.body = version_snap.body or article.body
    article.version += 1
    article.updated_by = user.id
    article.updated_at = datetime.now(timezone.utc)

    await db.commit()
    await db.refresh(article)

    breadcrumbs = await _get_breadcrumbs(db, article.section_id)
    creator = None
    if article.created_by:
        r = await db.execute(select(User).where(User.id == article.created_by))
        creator = r.scalar_one_or_none()
    return _article_to_public(article, breadcrumbs, creator, user)


# ── Комментарии ───────────────────────────────────────────────────────────────

@router.get("/articles/{article_id}/comments", response_model=KbCommentList, summary="Комментарии статьи")
async def list_comments(
    article_id: uuid.UUID,
    db: DbDep,
    user: CurrentUser,
    limit: int = Query(default=20, ge=1, le=100),
    offset: int = Query(default=0, ge=0),
) -> KbCommentList:
    await _get_article_or_404(db, article_id)

    count_r = await db.execute(
        select(func.count()).where(KbArticleComment.article_id == article_id)
    )
    total = count_r.scalar_one()

    result = await db.execute(
        select(KbArticleComment)
        .where(KbArticleComment.article_id == article_id)
        .order_by(KbArticleComment.created_at.asc())
        .limit(limit).offset(offset)
    )
    comments = result.scalars().all()

    user_ids = {c.author_id for c in comments if c.author_id}
    users_map: dict[uuid.UUID, User] = {}
    if user_ids:
        u_r = await db.execute(select(User).where(User.id.in_(user_ids)))
        for u in u_r.scalars():
            users_map[u.id] = u

    items = []
    for c in comments:
        author = users_map.get(c.author_id) if c.author_id else None
        items.append(KbCommentPublic(
            id=c.id,
            article_id=c.article_id,
            body=None if c.deleted_at else c.body,
            is_deleted=c.deleted_at is not None,
            created_at=c.created_at,
            updated_at=c.updated_at,
            author=KbUserRef(id=author.id, full_name=author.full_name, avatar_url=author.avatar_url) if author and not c.deleted_at else None,
        ))

    return KbCommentList(items=items, total=total)


@router.post("/articles/{article_id}/comments", status_code=status.HTTP_201_CREATED,
             response_model=KbCommentPublic, summary="Добавить комментарий")
async def create_comment(
    article_id: uuid.UUID,
    body: CreateCommentRequest,
    db: DbDep,
    user: CurrentUser,
) -> KbCommentPublic:
    await _get_article_or_404(db, article_id)

    comment = KbArticleComment(
        article_id=article_id,
        author_id=user.id,
        body=sanitize_html(body.body),
    )
    db.add(comment)
    await db.commit()
    await db.refresh(comment)
    return KbCommentPublic(
        id=comment.id,
        article_id=comment.article_id,
        body=comment.body,
        is_deleted=False,
        created_at=comment.created_at,
        updated_at=comment.updated_at,
        author=KbUserRef(id=user.id, full_name=user.full_name, avatar_url=user.avatar_url),
    )


@router.delete("/articles/{article_id}/comments/{comment_id}",
               status_code=status.HTTP_204_NO_CONTENT, summary="Удалить комментарий")
async def delete_comment(
    article_id: uuid.UUID,
    comment_id: uuid.UUID,
    db: DbDep,
    user: CurrentUser,
) -> None:
    result = await db.execute(
        select(KbArticleComment).where(
            KbArticleComment.id == comment_id,
            KbArticleComment.article_id == article_id,
        )
    )
    comment = result.scalar_one_or_none()
    if not comment:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Comment not found")
    if comment.deleted_at:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Already deleted")
    if user.role != "admin" and comment.author_id != user.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Insufficient permissions")

    comment.deleted_at = datetime.now(timezone.utc)
    await db.commit()


# ── Предложения правок ────────────────────────────────────────────────────────

@router.post("/articles/{article_id}/suggest", status_code=status.HTTP_202_ACCEPTED,
             summary="Предложить правку")
async def suggest_edit(
    article_id: uuid.UUID,
    body: CreateSuggestionRequest,
    db: DbDep,
    user: CurrentUser,
) -> dict:
    article = await _get_article_or_404(db, article_id)
    suggestion = KbSuggestion(
        article_id=article_id,
        author_id=user.id,
        body=body.body,
        comment=body.comment,
        status="pending",
    )
    db.add(suggestion)
    await db.commit()
    await db.refresh(suggestion)
    logger.info("kb.suggestion_created", article_id=str(article_id),
                suggestion_id=str(suggestion.id), user_id=str(user.id))
    return {"suggestion_id": str(suggestion.id), "message": "Правка отправлена на рассмотрение"}


@router.get("/articles/{article_id}/suggestions", summary="Список правок (editor+)")
async def list_suggestions(
    article_id: uuid.UUID,
    db: DbDep,
    user: EditorDep,
) -> dict:
    await _get_article_or_404(db, article_id)
    result = await db.execute(
        select(KbSuggestion).where(KbSuggestion.article_id == article_id).order_by(KbSuggestion.created_at.desc())
    )
    suggestions = result.scalars().all()
    user_ids = {s.author_id for s in suggestions if s.author_id}
    users_map: dict[uuid.UUID, User] = {}
    if user_ids:
        u_r = await db.execute(select(User).where(User.id.in_(user_ids)))
        for u in u_r.scalars():
            users_map[u.id] = u

    items = []
    for s in suggestions:
        author = users_map.get(s.author_id) if s.author_id else None
        items.append(KbSuggestionPublic(
            id=s.id, article_id=s.article_id, body=s.body, comment=s.comment,
            status=s.status, reviewed_at=s.reviewed_at, created_at=s.created_at,
            author=KbUserRef(id=author.id, full_name=author.full_name, avatar_url=author.avatar_url) if author else None,
        ))
    return {"items": items}


@router.post("/suggestions/{suggestion_id}/review", summary="Принять/отклонить правку (editor+)")
async def review_suggestion(
    suggestion_id: uuid.UUID,
    body: ReviewSuggestionRequest,
    db: DbDep,
    user: EditorDep,
) -> dict:
    result = await db.execute(select(KbSuggestion).where(KbSuggestion.id == suggestion_id))
    suggestion = result.scalar_one_or_none()
    if not suggestion:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Suggestion not found")
    if suggestion.status != "pending":
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Already reviewed")

    suggestion.status = "approved" if body.action == "approve" else "rejected"
    suggestion.reviewed_by = user.id
    suggestion.reviewed_at = datetime.now(timezone.utc)
    await db.commit()
    return {"status": suggestion.status}


# ── Обратная связь ────────────────────────────────────────────────────────────

@router.post("/articles/{article_id}/feedback", response_model=FeedbackStats,
             summary="Оценить статью (полезна/нет)")
async def submit_feedback(
    article_id: uuid.UUID,
    body: FeedbackRequest,
    db: DbDep,
    user: CurrentUser,
) -> FeedbackStats:
    await _get_article_or_404(db, article_id)

    existing = await db.execute(
        select(KbArticleFeedback).where(
            KbArticleFeedback.article_id == article_id,
            KbArticleFeedback.user_id == user.id,
        )
    )
    fb = existing.scalar_one_or_none()
    if fb:
        fb.is_helpful = body.is_helpful
    else:
        fb = KbArticleFeedback(article_id=article_id, user_id=user.id, is_helpful=body.is_helpful)
        db.add(fb)
    await db.commit()

    helpful_r = await db.execute(
        select(func.count()).where(KbArticleFeedback.article_id == article_id, KbArticleFeedback.is_helpful.is_(True))
    )
    not_helpful_r = await db.execute(
        select(func.count()).where(KbArticleFeedback.article_id == article_id, KbArticleFeedback.is_helpful.is_(False))
    )
    return FeedbackStats(
        helpful_count=helpful_r.scalar_one(),
        not_helpful_count=not_helpful_r.scalar_one(),
        user_feedback=body.is_helpful,
    )


# ── Экспорт ───────────────────────────────────────────────────────────────────

@router.post("/articles/{article_id}/export/pdf", summary="Экспорт статьи в PDF")
async def export_article_pdf(
    article_id: uuid.UUID,
    db: DbDep,
    user: CurrentUser,
    redis: RedisDep,
) -> Response:
    from app.core.pdf import render_pdf
    import markdown_it

    article = await _get_article_or_404(db, article_id)
    if article.status != "published" and user.role not in ("editor", "admin"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

    md = markdown_it.MarkdownIt()
    body_html = md.render(article.body)
    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
  body {{ font-family: Arial, sans-serif; margin: 0; padding: 0; color: #1a1a2e; }}
  h1 {{ font-size: 24px; margin-bottom: 8px; }}
  h2 {{ font-size: 18px; }}
  h3 {{ font-size: 16px; }}
  code {{ background: #f4f4f4; padding: 2px 4px; border-radius: 3px; font-size: 13px; }}
  pre {{ background: #f4f4f4; padding: 12px; border-radius: 4px; overflow-x: auto; }}
  blockquote {{ border-left: 3px solid #ccc; margin: 0; padding-left: 16px; color: #555; }}
  table {{ border-collapse: collapse; width: 100%; }}
  th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
</style></head><body>
<h1>{article.title}</h1>
{body_html}
</body></html>"""

    pdf_bytes = await render_pdf(html)
    safe_name = re.sub(r"[^\w\s-]", "", article.title)[:80].strip() or "article"
    filename = f"{safe_name}.pdf"
    encoded = filename.encode("utf-8").decode("latin-1", "replace")
    disposition = f"attachment; filename*=UTF-8''{filename}"
    await push_audit_event(
        redis,
        event_type="kb.article_exported_pdf",
        user_id=str(user.id),
        user_email=user.email,
        resource_type="kb_article",
        resource_id=str(article_id),
    )
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": disposition},
    )


@router.post("/articles/{article_id}/export/docx", summary="Экспорт статьи в DOCX")
async def export_article_docx(
    article_id: uuid.UUID,
    db: DbDep,
    user: CurrentUser,
    redis: RedisDep,
) -> Response:
    import io
    import markdown_it
    from docx import Document
    from docx.shared import Pt

    article = await _get_article_or_404(db, article_id)
    if article.status != "published" and user.role not in ("editor", "admin"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

    doc = Document()
    doc.add_heading(article.title, level=0)

    md = markdown_it.MarkdownIt()
    tokens = md.parse(article.body)
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "heading_open":
            level = int(tok.tag[1])
            content_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            heading_text = content_tok.content if content_tok else ""
            doc.add_heading(heading_text, level=min(level, 9))
            i += 3
        elif tok.type == "paragraph_open":
            content_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            para_text = content_tok.content if content_tok else ""
            doc.add_paragraph(para_text)
            i += 3
        elif tok.type == "fence":
            p = doc.add_paragraph()
            run = p.add_run(tok.content)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            i += 1
        elif tok.type == "bullet_list_open":
            i += 1
        elif tok.type in ("list_item_open",):
            i += 1
        elif tok.type == "inline" and i > 0 and tokens[i - 1].type == "list_item_open":
            doc.add_paragraph(tok.content, style="List Bullet")
            i += 1
        else:
            i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()

    safe_name = re.sub(r"[^\w\s-]", "", article.title)[:80].strip() or "article"
    filename = f"{safe_name}.docx"
    disposition = f"attachment; filename*=UTF-8''{filename}"
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    await push_audit_event(
        redis,
        event_type="kb.article_exported_docx",
        user_id=str(user.id),
        user_email=user.email,
        resource_type="kb_article",
        resource_id=str(article_id),
    )
    return Response(content=docx_bytes, media_type=mime, headers={"Content-Disposition": disposition})
