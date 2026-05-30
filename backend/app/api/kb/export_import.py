"""KB export (MD, ZIP, PDF, DOCX) and import endpoints."""

from __future__ import annotations

import io
import re
import uuid
import zipfile
from datetime import UTC, datetime
from pathlib import Path

from fastapi import APIRouter, HTTPException, Query, UploadFile, status
from fastapi.responses import Response, StreamingResponse
from sqlalchemy import select

from app.api.deps import CurrentUser, DbDep, RedisDep
from app.core.sanitize import sanitize_markdown
from app.core.system_config import load_system_settings
from app.core.text import slugify as _slugify_common
from app.models.kb import KbArticle, KbArticleTag, KbSection, KbTag
from app.models.user import User
from app.schemas.kb_extra import ImportReport
from app.services.audit import push_audit_event
from app.services.kb_acl import (
    batch_resolve_section_permissions,
    require_article_permission,
    require_section_permission,
    resolve_article_permission,
)

from ._common import _get_article_or_404, _rfc5987_filename
from ._frontmatter import (
    _build_frontmatter,
    _get_or_create_section_by_path,
    _get_section_path,
    _parse_frontmatter,
    _zip_section,
)

router = APIRouter(prefix="/kb", tags=["knowledge-base"])


def _kb_import_max_bytes() -> int:
    return load_system_settings().kb_import_max_size_mb * 1024 * 1024


@router.get("/articles/{article_id}/export/md")
async def export_article_md(
    article_id: uuid.UUID,
    db: DbDep,
    user: CurrentUser,
    redis: RedisDep,
) -> Response:
    article = await _get_article_or_404(db, article_id)
    await require_article_permission(user, article, "viewer", db, redis)

    author_res = (
        await db.execute(select(User.full_name).where(User.id == article.created_by))
        if article.created_by
        else None
    )
    author_name = author_res.scalar_one_or_none() if author_res else None
    section_path = await _get_section_path(db, article.section_id)
    frontmatter = _build_frontmatter(article, section_path, author_name)
    content = frontmatter + (article.body or "")

    safe_title = re.sub(r"[^\w\- ]", "", article.title)[:60].strip() or "article"
    filename = f"{safe_title}.md"
    cd = _rfc5987_filename(filename)
    await push_audit_event(
        redis,
        event_type="kb.article_exported_md",
        user_id=str(user.id),
        user_email=user.email,
        resource_type="kb_article",
        resource_id=str(article_id),
    )
    return Response(
        content=content.encode("utf-8"),
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": cd},
    )


@router.get("/sections/{section_id}/export/zip")
async def export_section_zip(
    section_id: uuid.UUID,
    db: DbDep,
    user: CurrentUser,
    redis: RedisDep,
) -> StreamingResponse:
    sec_res = await db.execute(select(KbSection).where(KbSection.id == section_id))
    section = sec_res.scalar_one_or_none()
    if not section:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Section not found")
    await require_section_permission(user, section, "viewer", db, redis)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        await _zip_section(zf, section, db, user, redis, prefix="")
    buf.seek(0)

    safe_title = re.sub(r"[^\w\- ]", "", section.title)[:40] or "section"
    filename = f"{safe_title}.zip"
    cd = _rfc5987_filename(filename)
    return StreamingResponse(
        iter([buf.read()]),
        media_type="application/zip",
        headers={"Content-Disposition": cd},
    )


@router.get("/export/vault.zip")
async def export_vault(
    db: DbDep,
    user: CurrentUser,
    redis: RedisDep,
) -> StreamingResponse:
    sec_res = await db.execute(
        select(KbSection).where(KbSection.parent_id.is_(None)).order_by(KbSection.sort_order)
    )
    root_sections = sec_res.scalars().all()

    root_perms = await batch_resolve_section_permissions(user, list(root_sections), db, redis)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for sec in root_sections:
            perm = root_perms.get(sec.id)
            if perm is None and user.role != "admin":
                continue
            await _zip_section(zf, sec, db, user, redis, prefix="")
    buf.seek(0)
    return StreamingResponse(
        iter([buf.read()]),
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=kb-vault.zip"},
    )


async def _import_single_article(
    db: DbDep,
    user: CurrentUser,
    redis: RedisDep,
    title: str,
    body: str,
    tags: list[str],
    section_path: str | None,
    strategy: str,
) -> str:
    section_id: uuid.UUID | None = None
    if section_path:
        section_id = await _get_or_create_section_by_path(db, section_path, user.id)

    existing_stmt = select(KbArticle).where(
        KbArticle.title == title, KbArticle.deleted_at.is_(None)
    )
    if section_id is None:
        existing_stmt = existing_stmt.where(KbArticle.section_id.is_(None))
    else:
        existing_stmt = existing_stmt.where(KbArticle.section_id == section_id)
    existing_res = await db.execute(existing_stmt)
    existing = existing_res.scalar_one_or_none()

    if existing:
        if strategy == "skip":
            return "skipped"
        elif strategy == "overwrite":
            await require_article_permission(user, existing, "editor", db, redis)
            existing.body = sanitize_markdown(body)
            existing.updated_at = datetime.now(UTC)
            existing.updated_by = user.id
            await db.flush()
            return "updated"
        else:
            title = f"{title} (импорт)"

    if section_id is not None:
        sec_res = await db.execute(
            select(KbSection).where(KbSection.id == section_id, KbSection.deleted_at.is_(None))
        )
        section = sec_res.scalar_one_or_none()
        if section is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Section not found")
        await require_section_permission(user, section, "editor", db, redis)

    article = KbArticle(
        title=title,
        body=sanitize_markdown(body),
        section_id=section_id,
        status="draft",
        created_by=user.id,
        updated_by=user.id,
    )
    db.add(article)
    await db.flush()

    for tag_name in tags[:20]:
        tag_slug = _slugify_common(tag_name, fallback="tag")
        t_res = await db.execute(select(KbTag).where(KbTag.slug == tag_slug))
        tag_obj = t_res.scalar_one_or_none()
        if not tag_obj:
            tag_obj = KbTag(name=tag_name.strip(), slug=tag_slug)
            db.add(tag_obj)
            await db.flush()
        db.add(KbArticleTag(article_id=article.id, tag_id=tag_obj.id))
        await db.flush()

    return "created"


@router.post("/articles/import", response_model=ImportReport, status_code=201)
async def import_article_md(
    file: UploadFile,
    db: DbDep,
    user: CurrentUser,
    redis: RedisDep,
    strategy: str = Query(default="skip", pattern="^(skip|overwrite|create_new)$"),
) -> ImportReport:
    _kb_import_max_mb = load_system_settings().kb_import_max_size_mb
    max_bytes = _kb_import_max_mb * 1024 * 1024

    if file.size is not None and file.size > max_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_CONTENT_TOO_LARGE,
            detail=f"File too large (max {_kb_import_max_mb} MB)",
        )

    content_bytes = await file.read(max_bytes + 1)
    if len(content_bytes) > max_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_CONTENT_TOO_LARGE,
            detail=f"File too large (max {_kb_import_max_mb} MB)",
        )
    try:
        content = content_bytes.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="File must be UTF-8 encoded"
        ) from exc

    fm, body = _parse_frontmatter(content)
    _stem = Path(file.filename).stem if file.filename else None
    title = fm.get("title") or _stem or "Untitled"
    tags: list[str] = fm.get("tags") or []
    section_path: str | None = fm.get("section")

    outcome = await _import_single_article(
        db, user, redis, title, body, tags, section_path, strategy
    )
    await db.commit()

    if outcome == "skipped":
        return ImportReport(created=0, updated=0, skipped=1, errors=[])
    elif outcome == "updated":
        return ImportReport(created=0, updated=1, skipped=0, errors=[])
    else:
        return ImportReport(created=1, updated=0, skipped=0, errors=[])


@router.post("/import/vault", response_model=ImportReport, status_code=201)
async def import_vault_zip(
    file: UploadFile,
    db: DbDep,
    user: CurrentUser,
    redis: RedisDep,
    strategy: str = Query(default="skip", pattern="^(skip|overwrite|create_new)$"),
) -> ImportReport:
    _max_bytes = _kb_import_max_bytes()
    if file.size and file.size > _max_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_CONTENT_TOO_LARGE, detail="Vault archive too large"
        )

    # Stream upload into memory in chunks, bailing out as soon as the limit is exceeded.
    # Do NOT trust file.size — it may be missing or spoofed.
    buf = io.BytesIO()
    received = 0
    chunk_size = 64 * 1024
    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        received += len(chunk)
        if received > _max_bytes:
            raise HTTPException(
                status_code=status.HTTP_413_CONTENT_TOO_LARGE,
                detail="Vault archive too large",
            )
        buf.write(chunk)
    buf.seek(0)

    report = ImportReport(created=0, updated=0, skipped=0, errors=[])

    try:
        with zipfile.ZipFile(buf) as zf:
            infolist = zf.infolist()
            # Limit total number of files in archive to prevent denial of service
            if len(infolist) > 1000:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Archive contains too many files (limit: 1000)",
                )

            # Limit total uncompressed size to prevent zip-bomb / memory exhaustion
            total_uncompressed_size = sum(info.file_size for info in infolist)
            if total_uncompressed_size > _max_bytes * 5:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Uncompressed archive size is too large (zip-bomb protection)",
                )

            md_files = []
            for name in zf.namelist():
                if not name.endswith(".md"):
                    continue
                # Reject absolute paths, path traversal, backslashes
                if name.startswith("/") or ".." in name or "\\" in name:
                    report.errors.append(
                        f"{name}: Invalid path (traversal, absolute or backslashes not allowed)"
                    )
                    continue
                # Reject special characters in name that might be dangerous
                if any(c in name for c in '\x00\r\n\t*:?|<>""'):
                    report.errors.append(f"{name}: Invalid characters in filename")
                    continue
                md_files.append(name)

            for md_path in md_files:
                try:
                    async with db.begin_nested():
                        raw = zf.read(md_path).decode("utf-8")
                        fm, body = _parse_frontmatter(raw)

                        parts = Path(md_path).parts
                        section_path_from_zip = (
                            "/" + "/".join(parts[:-1]) if len(parts) > 1 else None
                        )
                        title = fm.get("title") or Path(md_path).stem or "Untitled"
                        tags: list[str] = fm.get("tags") or []
                        section_path = fm.get("section") or section_path_from_zip

                        outcome = await _import_single_article(
                            db, user, redis, title, body, tags, section_path, strategy
                        )
                        if outcome == "skipped":
                            report.skipped += 1
                        elif outcome == "updated":
                            report.updated += 1
                        else:
                            report.created += 1
                except Exception as e:
                    report.errors.append(f"{md_path}: {e}")

        await db.commit()
    except zipfile.BadZipFile as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT, detail="Invalid ZIP file"
        ) from exc

    return report


@router.get("/articles/{article_id}/export/pdf", summary="Экспорт статьи в PDF")
async def export_article_pdf(
    article_id: uuid.UUID,
    db: DbDep,
    user: CurrentUser,
    redis: RedisDep,
) -> Response:
    import markdown_it

    from app.core.pdf import render_pdf

    article = await _get_article_or_404(db, article_id)
    pdf_perm = await resolve_article_permission(user, article, db, redis)
    if pdf_perm is None:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN, detail="Insufficient KB permissions"
        )
    if article.status != "published" and pdf_perm not in ("editor", "manager"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

    md = markdown_it.MarkdownIt()
    body_html = md.render(article.body)
    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
  body {{ font-family: Arial, sans-serif; margin: 0; padding: 0; color: #1a1a2e; }}
  h1 {{ font-size: 24px; margin-bottom: 8px; }}
  h2 {{ font-size: 18px; }}
  h3 {{ font-size: 16px; }}
  code {{ background: #f4f4f4; padding: 2px 4px; border-radius: 3px; font-size: 13px; }}
  pre {{ background: #f4f4f4; padding: 12px; border-radius: 4px; overflow-x: auto; }}
  blockquote {{ border-left: 3px solid #ccc; margin: 0; padding-left: 16px; color: #555; }}
  table {{ border-collapse: collapse; width: 100%; }}
  th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
</style></head><body>
<h1>{article.title}</h1>
{body_html}
</body></html>"""

    pdf_bytes = await render_pdf(html)
    safe_name = re.sub(r"[^\w\s-]", "", article.title)[:80].strip() or "article"
    filename = f"{safe_name}.pdf"
    disposition = _rfc5987_filename(filename)
    await push_audit_event(
        redis,
        event_type="kb.article_exported_pdf",
        user_id=str(user.id),
        user_email=user.email,
        resource_type="kb_article",
        resource_id=str(article_id),
    )
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": disposition},
    )


@router.get("/articles/{article_id}/export/docx", summary="Экспорт статьи в DOCX")
async def export_article_docx(
    article_id: uuid.UUID,
    db: DbDep,
    user: CurrentUser,
    redis: RedisDep,
) -> Response:
    import io as _io

    import markdown_it
    from docx import Document
    from docx.shared import Pt

    article = await _get_article_or_404(db, article_id)
    docx_perm = await resolve_article_permission(user, article, db, redis)
    if docx_perm is None:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN, detail="Insufficient KB permissions"
        )
    if article.status != "published" and docx_perm not in ("editor", "manager"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

    doc = Document()
    doc.add_heading(article.title, level=0)

    md = markdown_it.MarkdownIt()
    tokens = md.parse(article.body)
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "heading_open":
            level = int(tok.tag[1])
            content_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            heading_text = content_tok.content if content_tok else ""
            doc.add_heading(heading_text, level=min(level, 9))
            i += 3
        elif tok.type == "paragraph_open":
            content_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            para_text = content_tok.content if content_tok else ""
            doc.add_paragraph(para_text)
            i += 3
        elif tok.type == "fence":
            p = doc.add_paragraph()
            run = p.add_run(tok.content)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            i += 1
        elif tok.type == "bullet_list_open" or tok.type in ("list_item_open",):
            i += 1
        elif tok.type == "inline" and i > 0 and tokens[i - 1].type == "list_item_open":
            doc.add_paragraph(tok.content, style="List Bullet")
            i += 1
        else:
            i += 1

    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()

    safe_name = re.sub(r"[^\w\s-]", "", article.title)[:80].strip() or "article"
    filename = f"{safe_name}.docx"
    disposition = _rfc5987_filename(filename)
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    await push_audit_event(
        redis,
        event_type="kb.article_exported_docx",
        user_id=str(user.id),
        user_email=user.email,
        resource_type="kb_article",
        resource_id=str(article_id),
    )
    return Response(
        content=docx_bytes, media_type=mime, headers={"Content-Disposition": disposition}
    )
