"""DOCX rendering for KB article export (with inline media embedding)."""

from __future__ import annotations

import io
import uuid
from pathlib import Path
from typing import Any

from app.models.kb import KbArticle

from ._kb_media import KB_MEDIA_URL_RE, kb_media_path

_DOCX_NATIVE_FORMATS = {"PNG", "JPEG", "GIF", "BMP", "TIFF"}
_DOCX_MAX_WIDTH_IN = 6.0
_DOCX_ASSUMED_DPI = 96.0


def _resolve_media_path(src: str) -> Path | None:
    """Resolve a KB media URL from article body to a local file path."""
    match = KB_MEDIA_URL_RE.search(src)
    if not match:
        return None
    try:
        article_id = uuid.UUID(match.group(1))
    except ValueError:
        return None
    return kb_media_path(article_id, match.group(2))


def _embed_image(doc: Any, path: Path) -> bool:
    """Add an image to the document, normalising unsupported formats to PNG.

    Returns ``True`` on success. WebP (allowed for KB media but unsupported by
    python-docx) and other exotic formats are converted to PNG via Pillow.
    """
    from docx.shared import Inches
    from PIL import Image

    try:
        with Image.open(path) as im:
            fmt = (im.format or "").upper()
            width_px = im.width
            stream: str | io.BytesIO
            if fmt in _DOCX_NATIVE_FORMATS:
                stream = str(path)
            else:
                buf = io.BytesIO()
                im.convert("RGB").save(buf, format="PNG")
                buf.seek(0)
                stream = buf
    except (OSError, ValueError):
        return False

    width_in = width_px / _DOCX_ASSUMED_DPI
    try:
        if width_in > _DOCX_MAX_WIDTH_IN:
            doc.add_picture(stream, width=Inches(_DOCX_MAX_WIDTH_IN))
        else:
            doc.add_picture(stream)
    except Exception:
        return False
    return True


def _render_inline_paragraph(doc: Any, inline_tok: Any) -> None:
    """Render a markdown paragraph, embedding any inline KB images as pictures."""
    children = inline_tok.children if inline_tok else None
    if not children:
        if inline_tok and inline_tok.content:
            doc.add_paragraph(inline_tok.content)
        return

    images = [c for c in children if c.type == "image"]
    if not images:
        doc.add_paragraph(inline_tok.content)
        return

    text = "".join(c.content for c in children if c.type == "text").strip()
    if text:
        doc.add_paragraph(text)

    for img in images:
        src = img.attrGet("src") or ""
        path = _resolve_media_path(src)
        if path is not None and _embed_image(doc, path):
            continue
        alt = (img.content or "").strip()
        if alt:
            doc.add_paragraph(alt)


def render_article_docx(article: KbArticle) -> bytes:
    """Render a KB article to DOCX bytes, embedding inline media images."""
    import markdown_it
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(article.title, level=0)

    md = markdown_it.MarkdownIt()
    tokens = md.parse(article.body or "")
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "heading_open":
            level = int(tok.tag[1])
            content_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            heading_text = content_tok.content if content_tok else ""
            doc.add_heading(heading_text, level=min(level, 9))
            i += 3
        elif tok.type == "paragraph_open":
            content_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            _render_inline_paragraph(doc, content_tok)
            i += 3
        elif tok.type == "fence":
            p = doc.add_paragraph()
            run = p.add_run(tok.content)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            i += 1
        elif tok.type == "bullet_list_open" or tok.type in ("list_item_open",):
            i += 1
        elif tok.type == "inline" and i > 0 and tokens[i - 1].type == "list_item_open":
            doc.add_paragraph(tok.content, style="List Bullet")
            i += 1
        else:
            i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
